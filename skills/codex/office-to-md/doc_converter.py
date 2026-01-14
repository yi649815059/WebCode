"""
DOC格式转换器
使用LibreOffice将.doc文件转换为.docx格式，以支持跨平台解析
"""

import os
import subprocess
import logging
import tempfile
import shutil
import platform
from pathlib import Path
from typing import Optional, Tuple

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocConverter:
    """DOC到DOCX转换器"""
    
    def __init__(self, libreoffice_path: Optional[str] = None):
        """
        初始化转换器
        
        Args:
            libreoffice_path: LibreOffice可执行文件路径，如果为None则自动检测
        """
        self.libreoffice_path = libreoffice_path or self._find_libreoffice()
        if not self.libreoffice_path:
            raise RuntimeError("未找到LibreOffice，请确保已安装LibreOffice")
        
        logger.info(f"使用LibreOffice路径: {self.libreoffice_path}")
    
    def _find_libreoffice(self) -> Optional[str]:
        """自动查找LibreOffice可执行文件"""
        # Windows常见路径
        windows_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\Program Files\LibreOffice 7\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice 7\program\soffice.exe",
        ]
        
        # Linux/Mac常见路径
        unix_paths = [
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "/usr/local/bin/libreoffice",
            "/usr/local/bin/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
        
        # 根据操作系统选择路径
        if platform.system() == "Windows":
            search_paths = windows_paths
        else:
            search_paths = unix_paths
        
        # 查找第一个存在的路径
        for path in search_paths:
            if os.path.exists(path):
                return path
        
        # 尝试从环境变量PATH中查找
        try:
            result = subprocess.run(
                ["soffice", "--version"] if platform.system() != "Windows" else ["soffice.exe", "--version"],
                capture_output=True,
                timeout=5
            )
            if result.returncode == 0:
                return "soffice" if platform.system() != "Windows" else "soffice.exe"
        except:
            pass
        
        return None
    
    def convert_doc_to_docx(self, doc_path: str, output_dir: Optional[str] = None) -> Tuple[bool, str]:
        """
        将.doc文件转换为.docx格式
        
        Args:
            doc_path: 输入的.doc文件路径
            output_dir: 输出目录，如果为None则使用临时目录
            
        Returns:
            (成功标志, 转换后的docx文件路径或错误信息)
        """
        try:
            doc_path = Path(doc_path)
            if not doc_path.exists():
                return False, f"文件不存在: {doc_path}"
            
            if doc_path.suffix.lower() != '.doc':
                return False, f"文件格式不是.doc: {doc_path.suffix}"
            
            # 确定输出目录
            if output_dir is None:
                output_dir = tempfile.gettempdir()
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"开始转换: {doc_path.name}")
            logger.info(f"输出目录: {output_dir}")
            
            # 构建LibreOffice命令
            # --headless: 无界面模式
            # --convert-to docx: 转换为docx格式
            # --outdir: 输出目录
            
            # Linux环境优化：优先使用xvfb-run以支持无显示服务器环境
            use_xvfb = False
            if platform.system() != "Windows":
                # 设置环境变量以支持无头服务器
                os.environ['SAL_USE_VCLPLUGIN'] = 'svp'
                
                # 检查是否有xvfb-run（更稳定）
                if shutil.which("xvfb-run"):
                    use_xvfb = True
                    logger.info("检测到xvfb-run，使用虚拟显示服务器")
            
            if use_xvfb:
                cmd = [
                    "xvfb-run",
                    "-a",  # 自动选择显示编号
                    str(self.libreoffice_path),
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(output_dir),
                    str(doc_path)
                ]
            else:
                cmd = [
                    str(self.libreoffice_path),
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(output_dir),
                    str(doc_path)
                ]
            
            logger.info(f"执行命令: {' '.join(cmd)}")
            
            # 执行转换
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60  # 60秒超时
            )
            
            if result.returncode != 0:
                logger.error(f"转换失败，返回码: {result.returncode}")
                logger.error(f"错误输出: {result.stderr}")
                return False, f"转换失败: {result.stderr}"
            
            # 确定输出文件路径
            output_file = output_dir / f"{doc_path.stem}.docx"
            
            if not output_file.exists():
                return False, f"转换后的文件未生成: {output_file}"
            
            logger.info(f"转换成功: {output_file}")
            logger.info(f"文件大小: {output_file.stat().st_size} bytes")
            
            return True, str(output_file)
            
        except subprocess.TimeoutExpired:
            logger.error("转换超时")
            return False, "转换超时（超过60秒）"
        except Exception as e:
            logger.error(f"转换过程出错: {str(e)}")
            return False, f"转换出错: {str(e)}"
    
    def convert_and_cleanup(self, doc_path: str, keep_temp: bool = False) -> Tuple[bool, str]:
        """
        转换doc文件并可选择性清理临时文件
        
        Args:
            doc_path: 输入的.doc文件路径
            keep_temp: 是否保留临时转换的docx文件
            
        Returns:
            (成功标志, 转换后的docx文件路径或错误信息)
        """
        success, result_path = self.convert_doc_to_docx(doc_path)
        
        if not success:
            return False, result_path
        
        # 如果不保留临时文件，则在使用后自动清理（这个由调用方决定）
        if not keep_temp:
            logger.info(f"临时文件将在使用后清理: {result_path}")
        
        return True, result_path


def test_converter():
    """测试转换器功能"""
    print("=" * 60)
    print("DOC到DOCX转换器测试")
    print("=" * 60)
    
    try:
        # 创建转换器实例
        converter = DocConverter()
        print(f"✓ LibreOffice已找到: {converter.libreoffice_path}")
        print()
        
        # 测试文件路径（需要实际的.doc文件）
        test_doc_files = [
            "./test/test.doc",
            "./uploads/test.doc",
            "test.doc"
        ]
        
        converted = False
        for test_file in test_doc_files:
            if os.path.exists(test_file):
                print(f"正在测试转换文件: {test_file}")
                print("-" * 60)
                
                success, result = converter.convert_doc_to_docx(test_file)
                
                if success:
                    print(f"✓ 转换成功!")
                    print(f"  输出文件: {result}")
                    print(f"  文件大小: {Path(result).stat().st_size} bytes")
                    converted = True
                    
                    # 验证文件可以被打开
                    try:
                        import docx
                        doc = docx.Document(result)
                        print(f"  文档段落数: {len(doc.paragraphs)}")
                        print(f"  文档表格数: {len(doc.tables)}")
                        print("✓ 转换后的文件可以正常读取")
                    except Exception as e:
                        print(f"✗ 转换后的文件读取失败: {str(e)}")
                else:
                    print(f"✗ 转换失败: {result}")
                
                print()
                break
        
        if not converted:
            print("提示: 未找到测试文件，请创建一个test.doc文件进行测试")
            print("或者在以下位置放置.doc文件:")
            for path in test_doc_files:
                print(f"  - {path}")
        
    except RuntimeError as e:
        print(f"✗ 错误: {str(e)}")
        print()
        print("解决方案:")
        print("1. 下载并安装LibreOffice: https://www.libreoffice.org/download/")
        print("2. Windows用户: 安装后LibreOffice通常在 C:\\Program Files\\LibreOffice")
        print("3. Linux用户: sudo apt install libreoffice 或 sudo yum install libreoffice")
        print("4. Mac用户: brew install --cask libreoffice")
    
    print()
    print("=" * 60)


if __name__ == "__main__":
    test_converter()

